# -*- coding: utf-8 -*-
import cv2 as cv
import os
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
import numpy as np
import easyocr
import fitz
from docx import Document
from docx.shared import RGBColor, Cm, Inches
import zipfile

# Char ::= #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]

def valid_xml_char_ordinal(c):
  codepoint = ord(c)
  # conditions ordered by presumed frequency
  return (
    0x20 <= codepoint <= 0xD7FF or
    codepoint in (0x9, 0xA, 0xD) or
    0xE000 <= codepoint <= 0xFFFD or
    0x10000 <= codepoint <= 0x10FFFF
    )

def textFromPDF(doc):
  text = ''
  for page in doc:
    text = text + page.get_text()
  return text

def doImage(doc):
  images = convert_from_path(doc)
  for i in range(len(images)):
    images[i].save("page"+str(i)+".jpg", 'JPEG')
  
def deleteImage(pageNum):
  for i in range(pageNum):
    file_png = "page"+str(i)+".jpg"
    os.remove(file_png)

def textFromPDF_OCR_1(file_png):
  img = cv.imread(file_png)
  text = pytesseract.image_to_string(Image.open(file_png))
  return text

def text_EASYOCR(IMAGE_PATH):
  reader = easyocr.Reader(['en', 'ru'], gpu = True)
  result = reader.readtext(IMAGE_PATH)
  text = ''
  for detection in result:
    text = text + '\n' + detection[1]
  return text

def doGoodImage(IMAGE_PATH, pageNum):
  img = cv.imread(IMAGE_PATH)
  gray = cv.cvtColor(img, cv.COLOR_BGR2GRAY)
  se = cv.getStructuringElement(cv.MORPH_RECT , (8,8))
  bg = cv.morphologyEx(gray, cv.MORPH_DILATE, se)
  out_gray = cv.divide(gray, bg, scale = 255)
  out_binary = cv.threshold(out_gray, 0, 255, cv.THRESH_OTSU)[1]
  file_png = "page"+str(pageNum)+".jpg"
  cv.imwrite(file_png, out_gray)


def write2doc(document, text):
  with open(document, 'w+') as f:
    f.write(text)


def makeTable(document, text, ocr):
  table = document.add_table(rows=3, cols=2)
  table.cell(0,0).text = 'Text'
  table.cell(1,0).text = 'OCR'
  table.cell(2,0).text = 'Res'
  table.cell(0,1).paragraphs[0].add_run(text)
  table.cell(1,1).paragraphs[0].add_run(ocr)
  table.columns[0].width = Inches(1.0)
  table.columns[1].width = Inches(5.5)
  return table

def makeRESULT(lst_text, lst_ocr, table, length):
  count = 0
  for j in range(length):
    count = 0
    length_text, length_textOCR = len(lst_text[j]), len(lst_ocr[j])
    lengthT = min(length_text, length_textOCR)
    for i in range(lengthT):
      count = i
      run = table.cell(2,1).paragraphs[0].add_run(lst_ocr[j][i])
      if lst_ocr[j][i] == lst_text[j][i]:
        # print(lst_ocr[j][i] + " " + lst_text[j][i])
        run.font.color.rgb = RGBColor(0, 255, 0)
      else:
        run.font.color.rgb = RGBColor(255, 0, 0)
    if length_text >= length_textOCR:  
      while (count < length_textOCR - 1):
        run = table.cell(2,1).paragraphs[0].add_run(lst_ocr[j][count])
        run.font.color.rgb = RGBColor(255, 0, 0)
        count += 1
      while (count < length_text - 1):
        run = table.cell(2,1).paragraphs[0].add_run('*')
        run.font.color.rgb = RGBColor(255, 0, 0)
        count += 1
    else:
      while (count < length_textOCR - 1):
        count += 1
        run = table.cell(2,1).paragraphs[0].add_run(lst_ocr[j][count])
        run.font.color.rgb = RGBColor(255, 0, 0)
    run = table.cell(2,1).paragraphs[0].add_run(' ')

  while (length < len(lst_text)):
    run = table.cell(2,1).paragraphs[0].add_run('***')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run = table.cell(2,1).paragraphs[0].add_run(' ')
    length += 1

  while (length < len(lst_ocr)):
    run = table.cell(2,1).paragraphs[0].add_run(lst_ocr[length])
    run.font.color.rgb = RGBColor(255, 0, 0)
    run = table.cell(2,1).paragraphs[0].add_run(' ')
    length += 1



file = os.listdir('PDF')
PDF = 'PDF/' + file[1]
pdfDocument = fitz.open(PDF)
document = Document()
pageNum = pdfDocument.pageCount
textPDF = ''
textOCR = ''
textDiff = ''
documentPDF = "ready/textPDF.txt"
documentOCR = "ready/textOCR.txt"
documentRES = 'ready/RESULT.doc'


# textPDF = textFromPDF(pdfDocument)
doImage(PDF)
for i in range(pageNum):
  file_png = "page"+str(i)+".jpg"
  tmpTextOCR = text_EASYOCR(file_png)
  # tmpTextOCR = textFromPDF_OCR_1(file_png)
  tmpTextPdf = pdfDocument.loadPage(i).getText("text")
  tmpTextOCR = ''.join(c for c in tmpTextOCR if valid_xml_char_ordinal(c))
  tmpTextPdf = ''.join(c for c in tmpTextPdf if valid_xml_char_ordinal(c))
  textPDF = textPDF + tmpTextPdf
  textOCR = textOCR + tmpTextOCR
  lst_text = tmpTextPdf.split()
  lst_ocr = tmpTextOCR.split()
  length = min(len(lst_text), len(lst_ocr))
  table = makeTable(document, tmpTextPdf, tmpTextOCR)
  makeRESULT(lst_text, lst_ocr, table, length)


write2doc(documentOCR, textOCR)
write2doc(documentPDF, textPDF)
document.save(documentRES)
deleteImage(pageNum)
with zipfile.ZipFile('ready/text.zip','w') as zip:
  zip.write('ready/RESULT.doc')
  zip.write('ready/textOCR.txt')
  zip.write('ready/textPDF.txt')