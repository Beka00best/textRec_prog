from docx import Document
from docx.shared import RGBColor, Cm, Inches
document = Document()
p = document.add_paragraph()
# text = 'ТНЕ СВЕАТ ЕМСТ1$Н РОСОМЕМТ\n 1) Е15т${ Нет мВ зоте еп 1$В {ехё\n 2) бесопа Нет м/ИН зоте еуеп тоге 121158 {ехё. Ге те зрееК Нот ту ВеагЕ\n Топаоп\n ТБе сарНа!| оЁ СгеаЕ Втиат\n брееК\n Егот ту Веаг\n'
# ocr = 'THE GREAT ENGLISH DOCUMENT\n 1) Fisrst item with some english text\n 2) Second item with some even more inglish text. Let me speek from my heart\n London\n The capital of Great Britain\n Speek\n From my heart\n'
text = 'TH CBEATA EMCT1$H POCOMENT'
ocr = 'THE GREAT ENGLISH DOCUMENT'
lst_text = text.split()
lst_ocr = ocr.split()

# column = max(len(lst_text), len(lst_ocr)) 
length = min(len(lst_text), len(lst_ocr))
  
table = document.add_table(rows=3, cols=2)
for i in range(0, 3, 3):
  table.cell(i,0).text = 'Text'
  table.cell(i + 1,0).text = 'OCR'
  table.cell(i + 2,0).text = 'Res'
  

table.cell(0,1).paragraphs[0].add_run(text)
table.cell(1,1).paragraphs[0].add_run(ocr)

table.columns[0].width = Inches(1.0)
table.columns[1].width = Inches(5.5)

count = 0
for j in range(length):
  count = 0
  length_text, length_textOCR = len(lst_text[j]), len(lst_ocr[j])
  lengthT = min(length_text, length_textOCR)
  for i in range(lengthT):
    count = i
    run = table.cell(2,1).paragraphs[0].add_run(lst_ocr[j][i])
    if lst_ocr[j][i] == lst_text[j][i]:
      # print(lst_ocr[j][i] + " " + lst_text[j][i])
      run.font.color.rgb = RGBColor(0, 255, 0)
    else:
      run.font.color.rgb = RGBColor(255, 0, 0)
  if length_text >= length_textOCR:  
    while (count < length_textOCR - 1):
      run = table.cell(2,1).paragraphs[0].add_run(lst_ocr[j][count])
      run.font.color.rgb = RGBColor(255, 0, 0)
      count += 1
    while (count < length_text - 1):
      run = table.cell(2,1).paragraphs[0].add_run('*')
      run.font.color.rgb = RGBColor(255, 0, 0)
      count += 1
  else:
    while (count < length_textOCR - 1):
      count += 1
      run = table.cell(2,1).paragraphs[0].add_run(lst_ocr[j][count])
      run.font.color.rgb = RGBColor(255, 0, 0)
  run = table.cell(2,1).paragraphs[0].add_run(' ')

while (length < len(lst_text)):
  run = table.cell(2,1).paragraphs[0].add_run('***')
  run.font.color.rgb = RGBColor(255, 0, 0)
  run = table.cell(2,1).paragraphs[0].add_run(' ')
  length += 1

while (length < len(lst_ocr)):
  run = table.cell(2,1).paragraphs[0].add_run(lst_ocr[length])
  run.font.color.rgb = RGBColor(255, 0, 0)
  run = table.cell(2,1).paragraphs[0].add_run(' ')
  length += 1

document.save('demo1.docx')   

# for i in range(step):
#   run = p.add_run(ocr[i])
#   if ocr[i] == text[i]:
#     run.font.color.rgb = RGBColor(0, 255, 0)
#   else:
#     run.font.color.rgb = RGBColor(255, 0, 0)

# document.save('demo1.docx')


# from docx import Document
# from docx.shared import RGBColor, Cm, Inches

# document = Document()
# table = document.add_table(rows=3, cols=3)

# cell = table.cell(0,0).paragraphs[0].add_run('unc blue')
# cell.font.color.rgb = RGBColor(75, 156, 211)

# cell = table.cell(1,1).paragraphs[0].add_run('nc state red')
# cell.font.color.rgb = RGBColor(204, 0, 0)

# cell = table.cell(2,2).paragraphs[0].add_run('duke blue')
# cell.font.color.rgb = RGBColor(0, 26, 87)

# def set_col_widths(table):
#   widths = (Inches(1), Inches(2), Inches(1.5))
#   for row in table.rows:
#     for idx, width in enumerate(widths):
#       row.cells[idx].width = width

# set_col_widths(table)

# document.save('foo.doc')















# import docx
# from docx.shared import RGBColor, Cm, Inches

# doc = docx.Document()
# doc.add_heading('Name: ', level=1)

# table = doc.add_table(rows=4, cols=2)
# table.columns[0].width = Inches(1.0)
# table.autofit = True
# table.columns[1].width = Inches(5.5)
# # table.cell(0,0).width = Inches(1.0)     # 1.2 * 914400
# # table.cell(1,0).width = Inches(1.0)
# # table.cell(2,0).width = Inches(1.0)
# # table.cell(3,0).width = Inches(1.0)
# # table.cell(0,1).width = 4846320    # 5.3 * 914400 
# # table.cell(1,1).width = 4846320
# # table.cell(2,1).width = 4846320
# # table.cell(3,1).width = 4846320

# table.cell(0,0).text = 'Time Zone'
# table.cell(1,0).text = 'Link'
# table.cell(1,1).text = 'https://www.google.com/ghghghghghhghghghghghghghghghghghghghhghghghghghghh'
# table.cell(2,0).text = 'Website'
# table.cell(3,0).text = 'Facebook'

# doc.save('test.docx')
